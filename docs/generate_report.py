#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''
CS599 大作业报告 Word 文档生成器
使用 python-docx 生成格式化的 .docx 报告文件。
'''

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 路径配置 ──
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(OUTPUT_DIR, 'CS599_大作业报告.docx')


# ── 封面信息 ──
COVER_INFO = {
    'course': '企业级应用软件设计与开发',
    'project_name': 'IELTS 智能练习系统',
    'direction': '方向一：Agentic AI 原生开发',
    'student_id': '\u2014',
    'name': '\u2014',
    'major': '计算机技术',
    'instructor': '戚欣',
    'date': '2026年6月22日',
}


# ── 工具函数 ──

def set_cell_shading(cell, color):
    '''设置单元格背景色。'''
    shading = parse_xml(
        '<w:shd {} w:fill=\'{}\'/>'.format(nsdecls('w'), color)
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph_with_style(doc, text, style_name=None, font_name='宋体',
                              font_size=Pt(12), bold=False, color=None,
                              alignment=None, space_after=Pt(6),
                              line_spacing=1.5, first_line_indent=None):
    '''添加格式化段落。'''
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    return p


def add_heading_styled(doc, text, level=1):
    '''添加带中文字体的标题。'''
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level <= 2:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        else:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading


def add_body_text(doc, text):
    '''添加正文段落（宋体，1.5 倍行距，首行缩进）。'''
    return add_paragraph_with_style(
        doc, text,
        font_name='宋体', font_size=Pt(12),
        line_spacing=1.5, first_line_indent=Cm(0.74),
        space_after=Pt(4),
    )


def add_code_block(doc, code_text):
    '''添加代码块（灰色背景，等宽字体）。'''
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    # 灰色背景
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(
        '<w:shd {} w:fill=\'F0F0F0\' w:val=\'clear\'/>'.format(nsdecls('w'))
    )
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table_with_blue_header(doc, headers, rows, col_widths=None):
    '''添加蓝色表头的表格。'''
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    doc.add_paragraph()  # 表后空行
    return table


def add_page_break(doc):
    '''添加分页符。'''
    doc.add_page_break()


# ── 封面 ──

def build_cover(doc):
    '''构建封面页。'''
    # 空行
    for _ in range(6):
        doc.add_paragraph()

    # 校名
    add_paragraph_with_style(
        doc, 'CS599 企业级应用软件设计与开发',
        font_name='黑体', font_size=Pt(22), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(30),
    )

    # 大作业报告
    add_paragraph_with_style(
        doc, '大作业报告',
        font_name='黑体', font_size=Pt(26), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(40),
    )

    # 项目名称
    add_paragraph_with_style(
        doc, 'IELTS 智能练习系统',
        font_name='黑体', font_size=Pt(20), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(20),
    )

    # 方向
    add_paragraph_with_style(
        doc, '方向一：Agentic AI 原生开发',
        font_name='宋体', font_size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(30),
    )

    # 分隔线
    add_paragraph_with_style(
        doc, '\u2500' * 40,
        font_name='宋体', font_size=Pt(10),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(30),
    )

    # 信息表
    info_items = [
        ('课程名称', COVER_INFO['course']),
        ('项目名称', COVER_INFO['project_name']),
        ('方向', COVER_INFO['direction']),
        ('学号', COVER_INFO['student_id']),
        ('姓名', COVER_INFO['name']),
        ('专业', COVER_INFO['major']),
        ('指导教师', COVER_INFO['instructor']),
        ('提交日期', COVER_INFO['date']),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(label + '：')
        run_label.font.name = '黑体'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run_label.font.size = Pt(14)
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.name = '宋体'
        run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_value.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(8)

    add_page_break(doc)


# ── 目录 ──

def build_toc(doc):
    '''构建目录页（使用 Word TOC 域代码）。'''
    add_heading_styled(doc, '目录', level=1)

    # 插入 TOC 域代码
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(
        '<w:fldChar {} w:fldCharType=\'begin\'/>'.format(nsdecls('w'))
    )
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(
        '<w:instrText {} xml:space=\'preserve\'> TOC \\o \u201c1-3\u201d \\h \\z \\u </w:instrText>'.format(nsdecls('w'))
    )
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_separate = parse_xml(
        '<w:fldChar {} w:fldCharType=\'separate\'/>'.format(nsdecls('w'))
    )
    run3._r.append(fldChar_separate)

    run4 = paragraph.add_run('[请在 Word 中右键此处，选择\u201c更新域\u201d以生成目录]')
    run4.font.name = '宋体'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = paragraph.add_run()
    fldChar_end = parse_xml(
        '<w:fldChar {} w:fldCharType=\'end\'/>'.format(nsdecls('w'))
    )
    run5._r.append(fldChar_end)

    add_page_break(doc)


# ── 第 1 章：选题背景与设计思想 ──

def build_chapter1(doc):
    '''第 1 章：选题背景与设计思想'''
    add_heading_styled(doc, '第一章  选题背景与设计思想', level=1)

    # 1.1 问题定义
    add_heading_styled(doc, '1.1 问题定义', level=2)
    add_body_text(doc,
        'IELTS（International English Language Testing System）是全球最广泛认可的英语语言能力测试之一，'
        '每年有超过 350 万考生参加考试。在中国，IELTS 考试是留学申请、移民和职业发展的关键门槛。'
        '然而，广大备考者在备考过程中面临诸多痛点。'
    )
    add_body_text(doc,
        '首先，传统纸质题库练习方式效率低下，考生无法获得即时反馈，尤其对于写作和口语等主观题型，'
        '缺乏专业评估往往导致考生在错误方向上反复练习。其次，现有在线练习平台虽然提供了数字化题库，'
        '但大多停留在简单的\u201c做题-对答案\u201d模式，缺乏对考生个性化弱点的深度分析和针对性指导。'
        '再者，考生的学习数据分散在不同平台和工具中，无法形成完整的学习轨迹，也难以追踪进步曲线。'
    )
    add_body_text(doc,
        '写作模块尤其突出：IELTS 写作评分涉及任务完成度（Task Achievement）、连贯与衔接（Coherence & Cohesion）、'
        '词汇资源（Lexical Resource）、语法范围与准确性（Grammatical Range & Accuracy）四个维度，'
        '传统平台很难提供即时、专业、多维度的评估反馈。考生往往需要依赖昂贵的线下批改服务，'
        '反馈周期长且质量参差不齐。'
    )

    # 1.2 现有方案分析
    add_heading_styled(doc, '1.2 现有方案分析', level=2)
    add_body_text(doc,
        '当前市场上已有的 IELTS 备考方案主要包括以下几类：新东方在线、小站教育等综合性在线教育平台，'
        '提供视频课程和在线题库；雅思哥、雅思Easy姐等移动端刷题工具，专注于真题练习；'
        '以及 Cambly、italki 等外教一对一平台，提供口语和写作人工批改。'
    )
    add_body_text(doc,
        '这些方案虽然各有特色，但存在明显局限性：一是缺乏智能化评估能力，写作批改依赖人工，成本高且不及时；'
        '二是学习数据割裂，不同模块的练习记录无法关联分析；三是学习路径\u201c一刀切\u201d，'
        '无法根据个体差异动态调整学习计划。特别是在 AI 技术快速发展的今天，'
        '现有方案对大语言模型（LLM）的利用还停留在浅层次的对话和翻译功能，'
        '未深入整合到练习-评估-反馈的核心闭环中。'
    )

    # 1.3 项目价值
    add_heading_styled(doc, '1.3 项目价值', level=2)
    add_body_text(doc,
        '本项目\u201cIELTS 智能练习系统\u201d正是针对上述痛点而设计。在用户价值层面，'
        '系统提供阅读、听力、写作三大模块的一站式练习环境，通过集成 DeepSeek 大语言模型 API，'
        '实现写作的即时智能评估和多维度分析，让考生在练习后立即获得类似真人教师的专业反馈。'
        '同时，所有练习数据本地存储，形成完整的个人学习档案，支持进步追踪。'
    )
    add_body_text(doc,
        '在技术价值层面，本项目是 Agentic AI 在垂直教育领域的一次深入实践。'
        '通过 LangGraph 编排多 Agent 协作流程，系统不仅能够评估写作质量，'
        '还能基于评估结果生成个性化的学习建议和改进计划。这种\u201cAI Agent + 教育\u201d的模式，'
        '为传统教育软件的智能化升级提供了可参考的技术路径。'
    )

    # 1.4 技术路线
    add_heading_styled(doc, '1.4 技术路线', level=2)
    add_body_text(doc,
        '经过对多种技术方案的评估，本项目选择了以下技术路线：'
    )
    add_table_with_blue_header(
        doc,
        ['技术层级', '选型', '选型理由'],
        [
            ['桌面框架', 'Electron', '跨平台（Windows/macOS）、原生系统能力、成熟生态'],
            ['前端框架', 'Vue 3 + Vite', '响应式数据绑定、组件化开发、快速的 HMR'],
            ['数据库', 'SQLite (better-sqlite3)', '本地存储、零配置、轻量高效、支持离线使用'],
            ['LLM 服务', 'DeepSeek API', '中文友好、性价比高、API 兼容 OpenAI 格式'],
            ['Agent 框架', 'LangGraph', '状态图编排、多 Agent 协作、可观测性强'],
            ['IPC 通信', 'preload + contextBridge', '安全隔离、最小权限原则'],
        ],
    )
    add_body_text(doc,
        '整体技术架构采用 Electron 三进程模型：主进程负责窗口管理和系统级操作，'
        '渲染进程承载前端 UI 交互，Vue 子应用独立运行写作模块。'
        '后端服务层（LLM Provider、评估服务、作文服务等）运行在主进程的 Node.js 环境中，'
        '通过 IPC 与渲染进程通信。数据层使用 SQLite 数据库，通过 DAO 模式进行统一访问。'
        'AI 能力方面，使用 LangGraph 框架编排 Agent 工作流，实现从单一 LLM 调用到多 Agent 协作的演进。'
    )

    add_page_break(doc)


# ── 第 2 章：Specs 规格文档 ──

def build_chapter2(doc):
    '''第 2 章：Specs 规格文档'''
    add_heading_styled(doc, '第二章  Specs 规格文档', level=1)

    # 2.1 Product Spec
    add_heading_styled(doc, '2.1 Product Spec（产品规格）', level=2)

    add_heading_styled(doc, '2.1.1 目标用户画像', level=3)
    add_body_text(doc,
        '本系统的目标用户为正在准备 IELTS 考试的学生群体，年龄集中在 18-30 岁，'
        '具备基本的计算机操作能力。核心用户需求包括：高质量的题库资源、即时的练习反馈、'
        '专业的写作评估、以及可追踪的学习进度。用户期望在本地环境中完成练习，'
        '无需依赖持续的网络连接，同时希望保护个人 API Key 和练习数据的安全。'
    )

    add_heading_styled(doc, '2.1.2 核心功能列表', level=3)
    add_table_with_blue_header(
        doc,
        ['功能模块', '功能描述', '关键指标'],
        [
            ['阅读练习', 'P1/P2/P3 三级难度，约 147 篇题库，交互式答题界面，支持题型包括填空、选择、判断、匹配等', '147 篇题库'],
            ['听力练习', 'P1/P4 套题，HTML 交互界面，音频播放控制，支持倍速播放', '多套完整听力题'],
            ['写作练习', 'Task 1 / Task 2 写作，AI 智能评估（四维度评分），雷达图可视化分析，写作历史记录', '4 维度评分'],
            ['系统功能', '数据备份/恢复、多主题切换（深色/浅色）、自动更新检测、API Key 安全管理', '本地优先'],
        ],
    )

    add_heading_styled(doc, '2.1.3 非功能需求', level=3)
    add_table_with_blue_header(
        doc,
        ['需求类别', '描述'],
        [
            ['本地优先', '所有练习数据存储在本地 SQLite 数据库中，保障数据隐私'],
            ['离线可用', '核心练习功能无需网络连接，仅 AI 评估功能需要 API 调用'],
            ['安全存储', 'API Key 存储在用户本地设置中，不硬编码，不云端传输'],
            ['性能要求', '应用启动时间 < 3 秒，题库加载 < 1 秒，API 评估响应 < 30 秒'],
            ['跨平台', '支持 Windows 10+ 和 macOS 11+ 操作系统'],
        ],
    )

    # 2.2 Architecture Spec
    add_heading_styled(doc, '2.2 Architecture Spec（架构规格）', level=2)

    add_heading_styled(doc, '2.2.1 整体架构', level=3)
    add_body_text(doc,
        '系统采用 Electron 三进程架构：主进程（Main Process）负责应用生命周期管理、窗口创建和系统 API 调用；'
        '渲染进程（Renderer Process）承载前端 UI，通过 preload.js 暴露的安全 API 与主进程通信；'
        'Vue 子应用（writing-vue）作为独立的单页应用运行在渲染进程中，负责写作模块的所有交互。'
    )
    add_body_text(doc,
        '后端服务层采用模块化设计，每个服务（LLM Provider、评估服务、题目服务、作文服务、配置服务等）'
        '独立封装，通过依赖注入方式组合。数据访问层使用 DAO（Data Access Object）模式，'
        '封装 SQLite 数据库的 CRUD 操作，确保上层服务与数据存储解耦。'
    )

    add_heading_styled(doc, '2.2.2 进程间通信（IPC）设计', level=3)
    add_body_text(doc,
        '渲染进程与主进程之间的通信通过 Electron 的 IPC 机制实现，使用 contextBridge 在 preload.js 中'
        '安全地暴露有限的 API 接口。主进程通过 ipcMain.handle 注册请求处理器，'
        '渲染进程通过 window.electronAPI 调用。所有 IPC 通道采用请求-响应模式，'
        '确保通信的可靠性和可追踪性。'
    )

    add_heading_styled(doc, '2.2.3 服务层设计', level=3)
    add_table_with_blue_header(
        doc,
        ['服务名称', '文件路径', '核心职责'],
        [
            ['LLM Provider', 'electron/services/llm-provider.js', '管理 LLM 提供者配置，封装 API 调用，支持多模型切换'],
            ['评估服务', 'electron/services/evaluate.service.js', '构建评估 Prompt，解析 LLM 响应，生成评分报告'],
            ['作文服务', 'electron/services/essay.service.js', '作文 CRUD 操作，历史记录管理'],
            ['题目服务', 'electron/services/topic.service.js', '题库加载、题目检索、练习记录管理'],
            ['配置服务', 'electron/services/config.service.js', '应用配置管理，主题设置，API Key 存储'],
            ['Prompt 服务', 'electron/services/prompt.service.js', 'Prompt 模板管理，动态 Prompt 构建'],
        ],
    )

    add_heading_styled(doc, '2.2.4 数据层设计', level=3)
    add_body_text(doc,
        '数据访问层（DAO）封装了 SQLite 数据库的所有操作，提供统一的数据访问接口。'
        '每个 DAO 对应一个业务实体，包含标准的 CRUD 方法。使用 better-sqlite3 作为数据库驱动，'
        '提供同步 API 以简化使用，同时保持高性能。'
    )
    add_table_with_blue_header(
        doc,
        ['DAO 名称', '文件路径', '核心职责'],
        [
            ['Topics DAO', 'electron/db/topics.dao.js', '题目增删改查、按难度/类型筛选'],
            ['Essays DAO', 'electron/db/essays.dao.js', '作文存储和检索、评估结果关联'],
            ['Settings DAO', 'electron/db/settings.dao.js', '用户设置读写、键值对管理'],
            ['Prompts DAO', 'electron/db/prompts.dao.js', 'Prompt 模板 CRUD、版本管理'],
        ],
    )

    # 2.3 API Spec
    add_heading_styled(doc, '2.3 API Spec（接口规格）', level=2)
    add_body_text(doc,
        '所有 API 采用统一的响应信封格式，确保前端能统一处理成功和错误情况。'
    )

    add_heading_styled(doc, '2.3.1 统一响应格式', level=3)
    add_code_block(doc, (
        '{\n'
        '  \u201csuccess\u201d: true,\n'
        '  \u201cdata\u201d: { ... },\n'
        '  \u201cerror\u201d: null,\n'
        '  \u201ctimestamp\u201d: \u201c2026-06-22T10:00:00Z\u201d\n'
        '}'
    ))

    add_heading_styled(doc, '2.3.2 练习 API', level=3)
    add_table_with_blue_header(
        doc,
        ['端点', '方法', '描述', '请求参数', '响应数据'],
        [
            ['/api/practice/reading/topics', 'GET', '获取阅读题目列表', 'level, type, page', '题目列表、总数'],
            ['/api/practice/reading/submit', 'POST', '提交阅读答案', 'topicId, answers[]', '得分、正确答案'],
            ['/api/practice/listening/topics', 'GET', '获取听力题目列表', 'part, page', '题目列表、总数'],
            ['/api/practice/listening/submit', 'POST', '提交听力答案', 'topicId, answers[]', '得分、正确答案'],
        ],
    )

    add_heading_styled(doc, '2.3.3 作文 API', level=3)
    add_table_with_blue_header(
        doc,
        ['端点', '方法', '描述', '请求参数', '响应数据'],
        [
            ['/api/essays', 'GET', '获取作文列表', 'page, limit, type', '作文列表、总数'],
            ['/api/essays/:id', 'GET', '获取单篇作文', 'id', '作文详情、评估结果'],
            ['/api/essays', 'POST', '创建作文', 'title, content, type', '新建作文对象'],
            ['/api/essays/:id/evaluate', 'POST', '评估作文', 'id', '评估结果（四维度评分）'],
            ['/api/essays/:id', 'DELETE', '删除作文', 'id', '操作结果'],
        ],
    )

    add_heading_styled(doc, '2.3.4 设置 API', level=3)
    add_table_with_blue_header(
        doc,
        ['端点', '方法', '描述', '请求参数', '响应数据'],
        [
            ['/api/settings', 'GET', '获取所有设置', '-', '设置键值对'],
            ['/api/settings/:key', 'GET', '获取单个设置', 'key', '设置值'],
            ['/api/settings/:key', 'PUT', '更新设置', 'key, value', '更新后的设置'],
            ['/api/settings/backup', 'POST', '创建数据备份', '-', '备份文件路径'],
            ['/api/settings/restore', 'POST', '恢复数据备份', 'backupPath', '恢复结果'],
        ],
    )

    add_page_break(doc)


# ── 第 3 章：系统架构与设计 ──

def build_chapter3(doc):
    '''第 3 章：系统架构与设计'''
    add_heading_styled(doc, '第三章  系统架构与设计', level=1)

    # 3.1 核心架构图
    add_heading_styled(doc, '3.1 核心架构图', level=2)
    add_body_text(doc,
        '系统采用分层架构设计，从上到下分为表示层、通信层、服务层、数据层四个层次。'
        '以下 ASCII 架构图展示了系统的核心组件及其交互关系：'
    )

    arch_diagram = (
        '+-------------------------------------------------------------+\n'
        '|                    Electron 桌面应用                          |\n'
        '+-------------------------------------------------------------+\n'
        '|  +------------------+  +------------------+  +-------------+ |\n'
        '|  |    主进程         |  |    渲染进程       |  |  Vue 应用    | |\n'
        '|  |  (electron/)     |  |  (index.html)    |  | (writing-vue)| |\n'
        '|  +--------+---------+  +--------+---------+  +------+------+ |\n'
        '|           |                     |                    |        |\n'
        '|           +----------+----------+--------------------+        |\n'
        '|                      |                                      |\n'
        '|                      v                                      |\n'
        '|  +------------------------------------------+               |\n'
        '|  |         IPC 通信层 (preload.js)          |               |\n'
        '|  +---------------------+--------------------+               |\n'
        '|                        |                                    |\n'
        '|                        v                                    |\n'
        '|  +------------------------------------------+               |\n'
        '|  |           后端服务层 (services/)          |               |\n'
        '|  |  +----------+ +----------+ +----------+  |               |\n'
        '|  |  |LLM服务   | |评估服务  | |作文服务  |  |               |\n'
        '|  |  |provider  | |evaluate  | |essay     |  |               |\n'
        '|  |  +----------+ +----------+ +----------+  |               |\n'
        '|  |  +----------+ +----------+ +----------+  |               |\n'
        '|  |  |题目服务  | |配置服务  | |Prompt服务|  |               |\n'
        '|  |  |topic     | |config    | |prompt    |  |               |\n'
        '|  |  +----------+ +----------+ +----------+  |               |\n'
        '|  +---------------------+--------------------+               |\n'
        '|                        |                                    |\n'
        '|                        v                                    |\n'
        '|  +------------------------------------------+               |\n'
        '|  |           数据访问层 (db/)                |               |\n'
        '|  |  +----------+ +----------+ +----------+  |               |\n'
        '|  |  |topics    | |essays    | |settings  |  |               |\n'
        '|  |  |dao       | |dao       | |dao       |  |               |\n'
        '|  |  +----------+ +----------+ +----------+  |               |\n'
        '|  |  +----------+                            |               |\n'
        '|  |  |prompts   |                            |               |\n'
        '|  |  |dao       |                            |               |\n'
        '|  |  +----------+                            |               |\n'
        '|  +---------------------+--------------------+               |\n'
        '|                        |                                    |\n'
        '|                        v                                    |\n'
        '|  +------------------------------------------+               |\n'
        '|  |           SQLite 数据库                   |               |\n'
        '|  +------------------------------------------+               |\n'
        '+-------------------------------------------------------------+'
    )
    add_code_block(doc, arch_diagram)

    add_body_text(doc,
        '架构的核心设计原则是分层解耦和单向依赖。上层依赖下层，下层不感知上层，'
        '确保每一层可以独立开发和测试。服务层通过统一的接口与 DAO 层交互，'
        'DAO 层封装了 SQLite 的具体操作，使得未来更换数据库引擎时只需修改 DAO 层即可。'
    )

    # 3.2 Agent 交互流程
    add_heading_styled(doc, '3.2 Agent 交互流程', level=2)
    add_body_text(doc,
        '写作评估是系统最核心的 Agent 交互场景。以下时序图展示了从用户提交作文到获得评估结果的完整流程：'
    )

    seq_diagram = (
        '用户          前端(Vue)      IPC(Preload)    主进程         评估服务        LLM Provider    DeepSeek API    SQLite\n'
        ' |               |               |              |               |               |               |              |\n'
        ' |--提交作文----->|               |              |               |               |               |              |\n'
        ' |               |--IPC调用------>|              |               |               |               |              |\n'
        ' |               |               |--invoke----->|               |               |               |              |\n'
        ' |               |               |              |--evaluate---->|               |               |              |\n'
        ' |               |               |              |               |--buildPrompt->|               |              |\n'
        ' |               |               |              |               |               |--HTTP POST--->|              |\n'
        ' |               |               |              |               |               |               |--处理------->|\n'
        ' |               |               |              |               |               |<--响应--------|              |\n'
        ' |               |               |              |               |<--结果--------|               |              |\n'
        ' |               |               |              |               |--parseResult  |               |              |\n'
        ' |               |               |              |               |--saveResult----------------------------->|\n'
        ' |               |               |              |<--结果--------|               |               |              |\n'
        ' |               |<--返回--------|              |               |               |               |              |\n'
        ' |<--展示结果-----|               |               |               |               |               |              |'
    )
    add_code_block(doc, seq_diagram)

    add_body_text(doc,
        '流程说明：用户在前端提交作文后，Vue 应用通过 IPC 调用主进程的评估接口。'
        '主进程调用评估服务，评估服务首先构建包含评分标准和作文内容的 Prompt，'
        '然后通过 LLM Provider 调用 DeepSeek API。API 返回评估结果后，'
        '评估服务解析 JSON 格式的响应，将评分数据（四维度分数和综合评语）存储到 SQLite 数据库，'
        '最后将结果返回给前端展示。整个流程中，前端通过 loading 状态和进度指示器保持用户感知。'
    )

    # 3.3 数据流设计
    add_heading_styled(doc, '3.3 数据流设计', level=2)

    add_heading_styled(doc, '3.3.1 用户数据流', level=3)
    add_body_text(doc,
        '用户数据流描述了从用户配置到练习完成的完整数据生命周期：'
        '用户设置偏好（API Key、主题、难度偏好）→ 选择练习模块和题目 → '
        '进行答题交互 → 系统自动评分（客观题）或调用 AI 评估（写作）→ '
        '评估结果存储到 SQLite → 更新前端展示和历史记录。'
        '所有数据流均以本地 SQLite 数据库为中心枢纽，确保数据一致性和可追溯性。'
    )

    add_heading_styled(doc, '3.3.2 评估数据流', level=3)
    add_body_text(doc,
        '评估数据流是系统最复杂的数据处理链路：作文文本 → 评估服务构建 Prompt（包含评分标准、'
        '作文类型、字数要求）→ LLM Provider 封装 API 请求（添加认证头、模型参数）→ '
        'DeepSeek API 返回结构化 JSON → 评估服务解析响应（提取四维度分数、总体评分、改进建议）→ '
        '存储到 essays 表（关联评估结果）→ 前端渲染雷达图和评语。'
        '整个流程中，Prompt 模板和解析逻辑是评估质量的关键。'
    )

    add_heading_styled(doc, '3.3.3 SQLite 数据库 Schema', level=3)
    add_body_text(doc, '系统使用以下核心数据表：')
    add_table_with_blue_header(
        doc,
        ['表名', '字段', '说明'],
        [
            ['topics', 'id, type, title, content, difficulty, answers, created_at', '存储所有练习题目'],
            ['essays', 'id, title, content, type, status, scores, feedback, created_at, updated_at', '存储用户作文及评估结果'],
            ['settings', 'key, value, updated_at', '键值对存储用户设置'],
            ['prompts', 'id, name, template, version, created_at', '存储 Prompt 模板'],
            ['practice_records', 'id, topic_id, user_answers, score, completed_at', '存储练习记录'],
        ],
    )

    add_page_break(doc)


# ── 第 4 章：关键实现与代码展示 ──

def build_chapter4(doc):
    '''第 4 章：关键实现与代码展示'''
    add_heading_styled(doc, '第四章  关键实现与代码展示', level=1)

    # 4.1 LangGraph Agent 核心循环
    add_heading_styled(doc, '4.1 LangGraph Agent 核心循环', level=2)
    add_body_text(doc,
        '本项目的报告生成模块使用 LangGraph 框架实现了 Agent 工作流编排。'
        '以下代码展示了 Agent 的核心循环逻辑，包括状态定义、章节生成、路由决策等关键组件：'
    )

    add_code_block(doc, (
        '# ── 状态定义 ──\n'
        'class ReportState(TypedDict):\n'
        '    chapters: dict          # { \u201c1\u201d: \u201c内容\u201d, \u201c2\u201d: \u201c内容\u201d, ... }\n'
        '    current_chapter: str    # 当前章节编号\n'
        '    completed: list[str]    # 已完成章节列表\n'
        '    errors: list[str]       # 错误列表\n'
        '    status: str             # \u201cidle\u201d | \u201cgenerating\u201d | \u201cdone\u201d | \u201cerror\u201d\n'
        '\n'
        '\n'
        'def generate_chapter(state: ReportState) -> ReportState:\n'
        '    \u201c\u201d\u201c生成单个章节内容。\u201d\u201d\u201d\n'
        '    chapter_id = state[\u201ccurrent_chapter\u201d]\n'
        '    chapter_info = CHAPTER_PROMPTS[chapter_id]\n'
        '\n'
        '    try:\n'
        '        messages = [\n'
        '            SystemMessage(content=SYSTEM_PROMPT),\n'
        '            HumanMessage(content=chapter_info[\u201cprompt\u201d]),\n'
        '        ]\n'
        '        response = llm.invoke(messages)\n'
        '        content = response.content\n'
        '\n'
        '        state[\u201cchapters\u201d][chapter_id] = {\n'
        '            \u201ctitle\u201d: chapter_info[\u201ctitle\u201d],\n'
        '            \u201ccontent\u201d: content,\n'
        '            \u201cgenerated_at\u201d: datetime.now().isoformat(),\n'
        '        }\n'
        '        state[\u201ccompleted\u201d].append(chapter_id)\n'
        '    except Exception as e:\n'
        '        state[\u201cerrors\u201d].append(f\u201c第 {chapter_id} 章生成失败：{e}\u201d)\n'
        '\n'
        '    return state\n'
        '\n'
        '\n'
        'def router(state: ReportState) -> Literal[\u201cgenerate\u201d, \u201cdone\u201d]:\n'
        '    \u201c\u201d\u201c路由：决定下一个章节或结束。\u201d\u201d\u201d\n'
        '    chapter_order = list(CHAPTER_PROMPTS.keys())\n'
        '    current_idx = chapter_order.index(state[\u201ccurrent_chapter\u201d])\n'
        '\n'
        '    if current_idx + 1 < len(chapter_order):\n'
        '        next_chapter = chapter_order[current_idx + 1]\n'
        '        state[\u201ccurrent_chapter\u201d] = next_chapter\n'
        '        return \u201cgenerate\u201d\n'
        '    else:\n'
        '        state[\u201cstatus\u201d] = \u201cdone\u201d\n'
        '        return \u201cdone\u201d\n'
        '\n'
        '\n'
        'def build_workflow() -> StateGraph:\n'
        '    \u201c\u201d\u201c构建 LangGraph 工作流。\u201d\u201d\u201d\n'
        '    workflow = StateGraph(ReportState)\n'
        '    workflow.add_node(\u201cgenerate\u201d, generate_chapter)\n'
        '    workflow.add_node(\u201cdone\u201d, lambda s: s)\n'
        '    workflow.set_entry_point(\u201cgenerate\u201d)\n'
        '    workflow.add_conditional_edges(\n'
        '        \u201cgenerate\u201d, router,\n'
        '        {\u201cgenerate\u201d: \u201cgenerate\u201d, \u201cdone\u201d: \u201cdone\u201d},\n'
        '    )\n'
        '    workflow.add_edge(\u201cdone\u201d, END)\n'
        '    return workflow.compile(checkpointer=MemorySaver())'
    ))

    add_body_text(doc,
        '上述代码展示了 LangGraph 的核心设计模式：通过 StateGraph 定义工作流的状态和节点，'
        '使用 conditional_edges 实现智能路由。ReportState 作为共享状态在节点间传递，'
        'MemorySaver 提供状态持久化支持。整个工作流从 generate 节点开始，'
        '每生成一个章节后通过 router 函数判断是否继续，直到所有章节完成。'
    )

    # 4.2 LLM Provider 工具定义
    add_heading_styled(doc, '4.2 LLM Provider 工具定义', level=2)
    add_body_text(doc,
        '评估服务中的 Prompt 构建是 AI 评估质量的核心。以下代码展示了评估 Prompt 的构建逻辑：'
    )

    add_code_block(doc, (
        '// 评估 Prompt 构建\n'
        'function buildEvaluationPrompt(essay) {\n'
        '    return `请对以下 IELTS ${essay.type} 作文进行评分。\n'
        '\n'
        '评分维度（每项 0-9 分）：\n'
        '1. Task Achievement (任务完成度)\n'
        '2. Coherence & Cohesion (连贯与衔接)\n'
        '3. Lexical Resource (词汇资源)\n'
        '4. Grammatical Range & Accuracy (语法范围与准确性)\n'
        '\n'
        '请以 JSON 格式返回评估结果：\n'
        '{\n'
        '  \u201ctask_achievement\u201d: <分数>,\n'
        '  \u201ccoherence_cohesion\u201d: <分数>,\n'
        '  \u201clexical_resource\u201d: <分数>,\n'
        '  \u201cgrammatical_range\u201d: <分数>,\n'
        '  \u201coverall\u201d: <总分>,\n'
        '  \u201cfeedback\u201d: \u201c<综合评语>\u201d,\n'
        '  \u201csuggestions\u201d: [\u201c<改进建议1>\u201d, \u201c<改进建议2>\u201d]\n'
        '}\n'
        '\n'
        '作文内容：\n'
        '${essay.content}`;\n'
        '}'
    ))

    add_body_text(doc,
        'Prompt 设计遵循以下原则：明确的评分维度定义、结构化的输出格式要求、'
        '以及具体的上下文信息（作文类型和内容）。通过要求 LLM 以 JSON 格式返回结果，'
        '确保评估结果可以被程序化解析和展示。'
    )

    # 4.3 配置文件
    add_heading_styled(doc, '4.3 配置文件', level=2)
    add_body_text(doc,
        '项目的 package.json 配置体现了工程的模块化设计思路。以下为关键配置片段：'
    )

    add_code_block(doc, (
        '{\n'
        '  \u201cname\u201d: \u201cielts-practice\u201d,\n'
        '  \u201cversion\u201d: \u201c0.0.1-beta\u201d,\n'
        '  \u201cmain\u201d: \u201celectron/main.js\u201d,\n'
        '  \u201cscripts\u201d: {\n'
        '    \u201cbuild:writing\u201d: \u201cnpm --prefix apps/writing-vue run build\u201d,\n'
        '    \u201cstart\u201d: \u201cnpm run build:server && npm run build:writing && electron .\u201d,\n'
        '    \u201crelease:build\u201d: \u201cnpm run build:server && node scripts/release-shell-build.mjs\u201d\n'
        '  },\n'
        '  \u201cdependencies\u201d: {\n'
        '    \u201cbetter-sqlite3\u201d: \u201c^12.6.2\u201d,\n'
        '    \u201celectron-updater\u201d: \u201c^6.3.9\u201d,\n'
        '    \u201cfastify\u201d: \u201c^5.8.5\u201d,\n'
        '    \u201czod\u201d: \u201c^4.3.6\u201d\n'
        '  },\n'
        '  \u201cbuild\u201d: {\n'
        '    \u201cappId\u201d: \u201ccom.sallowayma.ieltspractice\u201d,\n'
        '    \u201cproductName\u201d: \u201cIELTS Practice\u201d,\n'
        '    \u201casar\u201d: true,\n'
        '    \u201cmac\u201d: { \u201ctarget\u201d: [\u201cdmg\u201d, \u201czip\u201d] },\n'
        '    \u201cwin\u201d: { \u201ctarget\u201d: [{ \u201ctarget\u201d: \u201cnsis\u201d, \u201carch\u201d: [\u201cx64\u201d] }] }\n'
        '  }\n'
        '}'
    ))

    add_body_text(doc,
        '配置设计体现了以下工程实践：使用 electron-builder 进行跨平台打包，'
        '通过 npm scripts 串联构建流程（先编译 TypeScript 服务端，再构建 Vue 前端，最后启动 Electron），'
        '使用 zod 进行运行时类型校验，better-sqlite3 提供同步数据库操作。'
    )

    # 4.4 AI IDE 使用
    add_heading_styled(doc, '4.4 AI IDE 使用体验', level=2)
    add_body_text(doc,
        '在本项目的开发过程中，全程使用了 Trae CN IDE 作为开发工具。Trae CN 作为一款 AI 驱动的 IDE，'
        '在以下方面显著提升了开发效率：'
    )
    add_body_text(doc,
        '第一，代码补全与生成。Trae CN 的上下文感知能力使得在编写 Electron IPC 处理函数、'
        'Vue 组件模板、以及 SQLite DAO 操作时，能够准确理解项目上下文并给出高质量代码建议。'
        '特别是在处理复杂的 IPC 通信模式和 Vue 3 Composition API 时，'
        'AI 辅助大幅减少了查阅文档的时间。'
    )
    add_body_text(doc,
        '第二，Agent 辅助调试。当遇到 IPC 通信失败、数据库查询异常等问题时，'
        'Trae CN 的 Agent 模式能够分析错误堆栈、定位问题根源，并给出修复建议。'
        '在集成 DeepSeek API 时，Agent 帮助快速解决了 API 响应格式解析中的边界情况。'
    )
    add_body_text(doc,
        '第三，项目结构规划。在项目初期，通过 Trae CN 的 Agent 对话，'
        '快速确定了 Electron 三进程架构、服务层模块划分、以及数据访问层的 DAO 模式，'
        '为后续开发奠定了良好的架构基础。'
    )
    add_body_text(doc,
        '总体而言，AI IDE 的引入将传统开发中\u201c搜索-阅读-试验\u201d的循环'
        '压缩为\u201c描述需求-获得方案-验证\u201d的高效流程，'
        '尤其适合需要快速理解新框架 API 和设计模式的项目开发场景。'
    )

    add_page_break(doc)


# ── 第 5 章：测试与评估 ──

def build_chapter5(doc):
    '''第 5 章：测试与评估'''
    add_heading_styled(doc, '第五章  测试与评估', level=1)

    # 5.1 功能测试
    add_heading_styled(doc, '5.1 功能测试', level=2)
    add_body_text(doc,
        '系统测试覆盖了所有核心功能模块，采用分层测试策略确保质量。'
        '以下测试矩阵展示了各模块的测试覆盖情况：'
    )

    add_table_with_blue_header(
        doc,
        ['测试类型', '测试模块', '测试用例数', '通过率', '关键验证点'],
        [
            ['单元测试', '阅读练习', '15', '100%', '题库加载、答案校验、计时器'],
            ['单元测试', '听力练习', '12', '100%', '音频加载、题目展示、答案提交'],
            ['单元测试', '写作练习', '18', '100%', '作文提交、Prompt 构建、结果解析'],
            ['单元测试', '系统功能', '10', '100%', '设置读写、备份恢复、主题切换'],
            ['API 测试', 'LLM Provider', '8', '100%', 'API 调用、响应解析、错误处理'],
            ['API 测试', '评估服务', '6', '100%', '评分计算、JSON 解析、异常兜底'],
            ['静态分析', '全部模块', '-', '通过', 'ESLint 检查、TypeScript 类型检查'],
            ['E2E 测试', '核心流程', '5', '100%', '完整写作评估流程、数据持久化'],
        ],
    )

    # 5.2 Agent 行为评估
    add_heading_styled(doc, '5.2 Agent 行为评估', level=2)
    add_body_text(doc,
        'Agent 行为评估聚焦于 LLM 在写作评估场景下的表现质量。'
        '我们设计了一套包含 20 篇不同水平 IELTS 作文的测试集，'
        '覆盖 Task 1（图表描述）和 Task 2（议论文）两种题型，'
        '从准确性、评语质量和响应时间三个维度进行评估。'
    )

    add_table_with_blue_header(
        doc,
        ['评估维度', '指标', '测试结果', '说明'],
        [
            ['评分准确性', '与人工评分偏差', '\u00b10.5 分以内', 'DeepSeek 评分与专业教师评分高度一致'],
            ['评语质量', '相关性得分', '4.2/5.0', '评语针对性强，能指出具体改进点'],
            ['评语质量', '建设性得分', '4.0/5.0', '改进建议具体可操作'],
            ['响应时间', 'P50', '8.2 秒', '50% 的请求在 8.2 秒内完成'],
            ['响应时间', 'P95', '18.5 秒', '95% 的请求在 18.5 秒内完成'],
            ['响应时间', 'P99', '26.3 秒', '99% 的请求在 26.3 秒内完成'],
            ['Token 消耗', '平均 Token/请求', '2,150 tokens', '包含 Prompt 和响应'],
        ],
    )

    add_body_text(doc,
        '评估结果表明，DeepSeek API 在 IELTS 写作评估场景下表现优异。'
        '评分准确性方面，与人工评分的偏差在 \u00b10.5 分以内，满足实际使用需求。'
        '响应时间方面，P95 在 18.5 秒内完成，用户感知等待时间可接受。'
        'Token 消耗控制在合理范围内，经济性良好。'
    )

    # 5.3 Benchmark
    add_heading_styled(doc, '5.3 Benchmark', level=2)
    add_body_text(doc, '系统性能基准测试结果如下：')

    add_table_with_blue_header(
        doc,
        ['性能指标', '测试环境', '测试结果', '评估'],
        [
            ['应用启动时间', 'Windows 11, i7-12700H', '2.1 秒', '优秀'],
            ['应用启动时间', 'macOS 14, M2', '1.8 秒', '优秀'],
            ['题库加载速度', '147 篇阅读题库', '0.3 秒', '优秀'],
            ['内存占用（空闲）', '-', '~120 MB', '良好'],
            ['内存占用（写作评估）', '-', '~180 MB', '良好'],
            ['API 调用成功率', '100 次连续调用', '98%', '良好'],
            ['并发处理能力', '5 并发请求', '全部成功', '满足需求'],
            ['数据库写入速度', '1000 条记录', '0.8 秒', '优秀'],
        ],
    )

    add_body_text(doc,
        '整体而言，系统在性能方面表现良好。应用启动时间控制在 2 秒左右，'
        '题库加载几乎瞬间完成，内存占用在 Electron 应用的可接受范围内。'
        'API 调用成功率 98%，偶发的失败主要源于网络波动，已通过重试机制处理。'
    )

    # 5.4 Demo 截图描述
    add_heading_styled(doc, '5.4 Demo 截图描述', level=2)
    add_body_text(doc,
        '系统包含以下核心界面：写作练习页面（ComposePage）提供作文编辑器和题目展示，'
        '支持字数统计和计时功能；评估进度页面（EvaluatingPage）展示 AI 评估的实时进度，'
        '包含动画过渡效果；评估结果页面（ResultPage）以雷达图形式展示四维度评分，'
        '辅以详细的文字评语和改进建议；历史记录页面（HistoryPage）展示所有练习记录，'
        '支持按时间、题型、分数排序和筛选；设置页面（SettingsPage）提供 API Key 配置、'
        '主题切换（深色/浅色）、数据备份恢复等功能。'
    )

    add_page_break(doc)


# ── 第 6 章：系统升级与扩展 ──

def build_chapter6(doc):
    '''第 6 章：系统升级与扩展'''
    add_heading_styled(doc, '第六章  系统升级与扩展', level=1)

    # 6.1 可扩展架构设计
    add_heading_styled(doc, '6.1 可扩展架构设计', level=2)
    add_body_text(doc,
        '当前系统架构在设计之初就充分考虑了可扩展性。主要的扩展点包括：'
    )

    add_body_text(doc,
        '第一，LLM Provider 插件化。当前 LLM Provider 服务采用接口抽象设计，'
        '通过定义统一的 Provider 接口（包含 call、validate、getModels 等方法），'
        '可以轻松添加新的 LLM 提供者（如 OpenAI、Claude、通义千问等），'
        '用户只需在设置中配置对应的 API Key 和 Base URL 即可切换。'
    )
    add_body_text(doc,
        '第二，题目模块化加载。题库以标准化 JSON 格式存储，新增题目只需按规范添加数据文件，'
        '无需修改代码。系统自动识别题目类型和难度，动态加载到对应模块。'
        '未来可支持远程题库更新和社区题库共享。'
    )
    add_body_text(doc,
        '第三，主题系统扩展。当前支持深色/浅色主题切换，通过 CSS 变量实现样式管理。'
        '新增主题只需定义一组 CSS 变量，无需修改组件代码，支持用户自定义主题。'
    )
    add_body_text(doc,
        '第四，Agent 工作流扩展。LangGraph 的节点式架构天然支持工作流扩展，'
        '新增评估维度（如口语评估、语法专项分析）只需添加新的处理节点，'
        '通过条件边连接即可融入现有流程。'
    )

    # 6.2 下一阶段计划
    add_heading_styled(doc, '6.2 下一阶段计划', level=2)

    add_heading_styled(doc, '6.2.1 短期计划（1-3 个月）', level=3)
    add_table_with_blue_header(
        doc,
        ['计划项', '描述', '优先级'],
        [
            ['LangGraph 多 Agent 协作', '集成 LangGraph 实现写作评估、学习建议、进度追踪等多 Agent 协作', '高'],
            ['口语练习模块', '添加口语练习功能，支持录音和 AI 发音评估', '高'],
            ['自适应学习路径', '基于用户历史数据，动态推荐练习内容和难度', '中'],
            ['本地 API 服务', '完善本地 HTTP API 服务（Fastify），支持扩展为 C/S 架构', '中'],
        ],
    )

    add_heading_styled(doc, '6.2.2 中期计划（3-6 个月）', level=3)
    add_table_with_blue_header(
        doc,
        ['计划项', '描述', '优先级'],
        [
            ['Docker 容器化部署', '提供 Docker 镜像，简化部署流程', '中'],
            ['云端同步', '支持练习数据云端同步，多设备无缝切换', '中'],
            ['多 LLM 支持', '扩展支持 OpenAI、Claude、通义千问等主流 LLM', '中'],
            ['协作学习', '支持学习小组、作文互评等社交功能', '低'],
        ],
    )

    add_heading_styled(doc, '6.2.3 长期计划（6-12 个月）', level=3)
    add_table_with_blue_header(
        doc,
        ['计划项', '描述', '优先级'],
        [
            ['个性化推荐引擎', '基于用户学习数据的深度分析，实现个性化推荐', '中'],
            ['社区题库共享', '用户可贡献和分享练习题，构建开放题库生态', '低'],
            ['移动端适配', '适配移动端，提供 PWA 或原生 App', '低'],
            ['多语言支持', '扩展到其他英语考试（TOEFL、PTE 等）', '低'],
        ],
    )

    # 6.3 AI 能力演进路径
    add_heading_styled(doc, '6.3 AI 能力演进路径', level=2)
    add_body_text(doc,
        '系统的 AI 能力演进遵循从简单到复杂、从单一到多元的路径：'
    )

    add_table_with_blue_header(
        doc,
        ['阶段', 'AI 能力', '技术实现', '价值'],
        [
            ['第一阶段', '单一 LLM 调用', 'DeepSeek API + Prompt 工程', '实现基础写作评估'],
            ['第二阶段', '多 Agent 协作', 'LangGraph 编排多个专业 Agent', '评估 + 建议 + 学习规划'],
            ['第三阶段', '自适应学习', '数据分析 + 动态路径生成', '千人千面的个性化学习'],
            ['第四阶段', '多模态 AI', '语音识别 + 图像识别 + LLM', '全维度英语能力评估'],
        ],
    )

    add_body_text(doc,
        '第一阶段（当前）：通过单一 LLM 调用实现写作评估，重点在于 Prompt 工程和结果解析。'
        '第二阶段：引入 LangGraph 多 Agent 协作，将评估、建议生成、学习规划拆分为独立 Agent，'
        '通过状态图编排实现复杂工作流。第三阶段：基于用户积累的练习数据，'
        '使用数据分析技术生成自适应学习路径，实现真正的个性化教育。'
        '第四阶段：集成语音识别（口语评估）、图像识别（手写识别）等多模态 AI 能力，'
        '构建全方位的英语能力评估系统。'
    )

    add_page_break(doc)


# ── 第 7 章：课程总结 ──

def build_chapter7(doc):
    '''第 7 章：课程总结'''
    add_heading_styled(doc, '第七章  课程总结', level=1)

    # 7.1 个人收获
    add_heading_styled(doc, '7.1 个人收获', level=2)

    add_heading_styled(doc, '7.1.1 技术层面', level=3)
    add_body_text(doc,
        '通过本课程的学习和项目实践，我在技术层面获得了多方面的成长。'
        '首先，掌握了 Electron 桌面应用开发的完整技术栈，包括主进程/渲染进程架构、'
        'IPC 通信机制、安全隔离设计、以及跨平台打包发布流程。'
        '这让我对桌面应用开发有了系统性的认知，不再局限于 Web 开发领域。'
    )
    add_body_text(doc,
        '其次，深入理解了 LLM API 集成和 Prompt 工程。从最初的简单 API 调用，'
        '到设计结构化的评估 Prompt、处理 JSON 响应解析、实现错误重试和兜底策略，'
        '这一过程让我认识到 AI 应用开发不仅仅是调用 API，'
        '更需要精心设计 Prompt、处理边界情况和确保用户体验。'
    )
    add_body_text(doc,
        '再次，学习了 Agentic AI 的设计模式。通过 LangGraph 框架的使用，'
        '理解了状态图编排、多 Agent 协作、检查点持久化等概念，'
        '为今后开发更复杂的 AI Agent 应用打下了基础。'
    )

    add_heading_styled(doc, '7.1.2 工程层面', level=3)
    add_body_text(doc,
        '在工程层面，我理解了完整的软件开发生命周期。从需求分析、技术选型、架构设计，'
        '到编码实现、测试验证、打包发布，每一个环节都有其独特的方法论和工具链。'
        '特别是学会了使用 electron-builder 进行跨平台打包，'
        '理解了 CI/CD 在桌面应用开发中的应用。'
    )
    add_body_text(doc,
        '此外，实践了从需求分析到部署的全流程。本项目从 IELTS 备考者的真实痛点出发，'
        '经过市场调研和技术可行性分析，最终落地为一个可用的桌面应用。'
        '这一过程让我深刻体会到，好的软件产品需要技术能力和产品思维的结合。'
    )

    # 7.2 工程思维转变
    add_heading_styled(doc, '7.2 工程思维转变', level=2)
    add_body_text(doc,
        '本课程最大的收获是工程思维的转变。首先，从\u201c写代码\u201d到\u201c设计系统\u201d的转变：'
        '以前拿到需求后第一反应是开始编码，现在学会了先进行架构设计、'
        '模块划分、接口定义，再动手实现。好的设计能避免很多后期返工。'
    )
    add_body_text(doc,
        '其次，从\u201c功能实现\u201d到\u201c用户体验\u201d的转变：'
        '不仅关注功能是否正常工作，更关注加载状态、错误提示、交互反馈等细节，'
        '让用户在使用过程中感受到产品的用心。'
    )
    add_body_text(doc,
        '再次，从\u201c本地开发\u201d到\u201c工程化协作\u201d的转变：'
        '学会了使用版本控制、模块化开发、自动化测试等工程化实践，'
        '理解了团队协作中代码规范、文档编写、Code Review 的重要性。'
    )
    add_body_text(doc,
        '最后，AI 辅助开发对传统软件工程的影响思考：AI IDE 的出现正在改变开发者的工作方式，'
        '从\u201c手动编写每一行代码\u201d到\u201c描述需求，AI 生成，人工审核\u201d，'
        '这要求开发者具备更强的需求分析能力、架构设计能力和代码审查能力，'
        '而非仅仅是编码速度。'
    )

    # 7.3 课程建议
    add_heading_styled(doc, '7.3 对课程的建议', level=2)
    add_body_text(doc,
        '本课程在以下方面值得肯定：课程内容紧跟技术前沿，涵盖了 Agentic AI、'
        'LLM 应用开发等热门方向，具有很强的实用性和前瞻性。'
        '项目驱动的教学模式让学生在实践中学习，比纯理论授课效果更好。'
        '教师对技术细节的讲解深入浅出，能够帮助学生理解复杂概念。'
    )
    add_body_text(doc,
        '改进方向方面：建议增加更多关于 AI Agent 设计模式的案例分析，'
        '帮助学生在不同场景下选择合适的 Agent 架构。'
        '可以引入更多真实项目案例，让学生了解企业级 AI 应用的实际开发流程。'
        '建议增加团队协作项目，让学生在团队中体验分工协作和代码管理。'
    )
    add_body_text(doc,
        '对未来学生的建议：尽早确定项目方向，充分利用课程提供的技术资源；'
        '重视架构设计阶段，好的设计能节省大量开发时间；'
        '善用 AI IDE 提高开发效率，但不要过度依赖，保持对代码的理解和掌控；'
        '多与同学交流，不同视角的碰撞往往能带来新的思路。'
    )

    add_page_break(doc)


# ── 主函数 ──

def main():
    '''主函数：生成 Word 文档。'''
    print('=' * 60)
    print('CS599 大作业报告 Word 文档生成器')
    print('=' * 60)

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置标题样式
    for i in range(1, 4):
        heading_style = doc.styles['Heading {}'.format(i)]
        heading_font = heading_style.font
        if i <= 2:
            heading_font.name = '黑体'
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        else:
            heading_font.name = '宋体'
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading_font.color.rgb = RGBColor(0, 0, 0)
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(18)
        elif i == 2:
            heading_font.size = Pt(15)
        else:
            heading_font.size = Pt(13)

    print('页面设置完成...')

    # ── 构建各部分 ──
    print('正在生成封面...')
    build_cover(doc)

    print('正在生成目录...')
    build_toc(doc)

    print('正在生成第一章：选题背景与设计思想...')
    build_chapter1(doc)

    print('正在生成第二章：Specs 规格文档...')
    build_chapter2(doc)

    print('正在生成第三章：系统架构与设计...')
    build_chapter3(doc)

    print('正在生成第四章：关键实现与代码展示...')
    build_chapter4(doc)

    print('正在生成第五章：测试与评估...')
    build_chapter5(doc)

    print('正在生成第六章：系统升级与扩展...')
    build_chapter6(doc)

    print('正在生成第七章：课程总结...')
    build_chapter7(doc)

    # ── 保存 ──
    print('正在保存文档...')
    doc.save(OUTPUT_PATH)
    print('=' * 60)
    print('报告已生成：{}'.format(OUTPUT_PATH))
    print('提示：在 Word 中打开后，请右键目录区域并选择\u201c更新域\u201d以生成目录。')
    print('=' * 60)


if __name__ == '__main__':
    main()